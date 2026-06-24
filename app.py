import json
import os
import re
from html import escape
from io import BytesIO
from typing import Any

import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph as DocxParagraph
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer


load_dotenv()

st.set_page_config(
    page_title="AI Internship Hunter",
    page_icon=":briefcase:",
    layout="centered",
)


SYSTEM_PROMPT = """
You are an expert career coach for students applying to internships.
Analyze how well a resume matches a job description.

Return only valid JSON with this exact shape:
{
  "match_score": 78,
  "strong_matches": ["Python", "SQL", "IT support"],
  "missing_keywords": ["AWS", "React", "Agile"],
  "resume_suggestions": [
    "Rewrite the project section to show measurable outcomes.",
    "Add a tools line that includes SQL and ticketing systems."
  ],
  "tailored_bullets": [
    "Built Python scripts to automate weekly reporting, reducing manual work by 30%.",
    "Used SQL to investigate support issues and summarize findings for non-technical users."
  ]
}

Rules:
- match_score must be an integer from 0 to 100.
- strong_matches and missing_keywords should be concise keywords or phrases.
- resume_suggestions should be specific, practical, and honest.
- tailored_bullets must not invent fake experience. Reframe only what is supported by the resume.
- If the resume lacks evidence for a requested skill, suggest a truthful learning/project bullet instead of pretending.
"""


COMMON_KEYWORDS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "sql",
    "r",
    "html",
    "css",
    "react",
    "node",
    "aws",
    "azure",
    "gcp",
    "docker",
    "git",
    "github",
    "linux",
    "excel",
    "tableau",
    "power bi",
    "pandas",
    "numpy",
    "tensorflow",
    "pytorch",
    "machine learning",
    "data analysis",
    "data visualization",
    "api",
    "rest",
    "agile",
    "scrum",
    "communication",
    "collaboration",
    "troubleshooting",
    "customer support",
    "it support",
    "documentation",
    "research",
    "statistics",
    "automation",
]


def get_api_key() -> str:
    try:
        return st.secrets.get("OPENAI_API_KEY", "") or os.getenv("OPENAI_API_KEY", "")
    except Exception:
        return os.getenv("OPENAI_API_KEY", "")


def get_model() -> str:
    return os.getenv("OPENAI_MODEL", "gpt-4.1-mini")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower())


def extract_keywords(text: str) -> list[str]:
    normalized = normalize_text(text)
    found = []
    for keyword in COMMON_KEYWORDS:
        pattern = r"\b" + re.escape(keyword) + r"\b"
        if re.search(pattern, normalized):
            found.append(keyword)
    return found


def title_keyword(keyword: str) -> str:
    special_cases = {
        "api": "API",
        "aws": "AWS",
        "azure": "Azure",
        "css": "CSS",
        "gcp": "GCP",
        "git": "Git",
        "github": "GitHub",
        "html": "HTML",
        "it support": "IT support",
        "javascript": "JavaScript",
        "linux": "Linux",
        "node": "Node.js",
        "power bi": "Power BI",
        "rest": "REST",
        "sql": "SQL",
    }
    return special_cases.get(keyword, keyword.title())


def demo_analyze_resume(resume: str, job_description: str) -> dict[str, Any]:
    resume_keywords = set(extract_keywords(resume))
    job_keywords = extract_keywords(job_description)
    job_keyword_set = set(job_keywords)

    strong_matches = sorted(resume_keywords & job_keyword_set)
    missing_keywords = sorted(job_keyword_set - resume_keywords)

    if job_keywords:
        keyword_score = len(strong_matches) / len(job_keyword_set)
    else:
        keyword_score = 0

    resume_length_score = min(len(resume.split()) / 450, 1)
    score = round((keyword_score * 75) + (resume_length_score * 25))
    score = max(15, min(score, 95))

    suggestions = [
        "Move the most relevant projects or experience closer to the top of the resume.",
        "Add measurable impact where possible, such as time saved, users supported, accuracy improved, or reports automated.",
        "Mirror important wording from the job description, but only when it honestly matches your experience.",
    ]

    if missing_keywords:
        top_missing = ", ".join(title_keyword(keyword) for keyword in missing_keywords[:4])
        suggestions.insert(0, f"Add honest evidence for these missing keywords if you have it: {top_missing}.")

    top_matches = strong_matches[:3] or list(resume_keywords)[:3]
    match_phrase = ", ".join(title_keyword(keyword) for keyword in top_matches) or "relevant technical skills"
    tailored_bullets = [
        f"Applied {match_phrase} to complete technical tasks and support project goals.",
        "Collaborated with teammates and stakeholders to clarify requirements, document work, and deliver usable results.",
        "Improved resume alignment by emphasizing tools, outcomes, and responsibilities that appear in the target internship description.",
    ]

    if missing_keywords:
        tailored_bullets.append(
            f"Built or planned a small project to develop exposure to {title_keyword(missing_keywords[0])} in a practical context."
        )

    return {
        "match_score": score,
        "strong_matches": [title_keyword(keyword) for keyword in strong_matches[:8]],
        "missing_keywords": [title_keyword(keyword) for keyword in missing_keywords[:8]],
        "resume_suggestions": suggestions,
        "tailored_bullets": tailored_bullets,
    }


def build_demo_tailored_resume(resume: str, job_description: str, result: dict[str, Any]) -> str:
    strong_matches = result.get("strong_matches", [])
    skills_line = ", ".join(strong_matches) if strong_matches else "Relevant technical and collaboration skills"
    sections = parse_resume_sections(resume)
    contact = sections["CONTACT"]

    output_sections = []
    if contact:
        output_sections.append(contact)

    ordered_headings = [
        "EDUCATION",
        "INTERNSHIP EXPERIENCE",
        "WORK EXPERIENCE",
        "EXPERIENCE",
        "PROJECT EXPERIENCE",
        "PROJECTS",
        "TECHNICAL SKILLS",
        "SKILLS",
    ]
    for heading in ordered_headings:
        if heading in sections and sections[heading]:
            if heading in {"TECHNICAL SKILLS", "SKILLS"} and strong_matches:
                section_text = sections[heading]
                if "targeted skills" not in section_text.lower():
                    section_text += f"\nTargeted Skills: {skills_line}"
                output_sections.append(section_text)
            else:
                output_sections.append(sections[heading])

    used_sections = {"CONTACT", *ordered_headings}
    for heading, text in sections.items():
        if heading not in used_sections and text:
            output_sections.append(text)

    return "\n\n".join(output_sections).strip()


def parse_resume_sections(resume: str) -> dict[str, str]:
    known_headings = {
        "EDUCATION",
        "INTERNSHIP EXPERIENCE",
        "EXPERIENCE",
        "WORK EXPERIENCE",
        "PROJECT EXPERIENCE",
        "PROJECTS",
        "TECHNICAL SKILLS",
        "SKILLS",
        "LANGUAGES & SKILLS",
        "CERTIFICATIONS",
        "LEADERSHIP",
        "ACTIVITIES",
        "AWARDS",
    }
    sections: dict[str, list[str]] = {"CONTACT": []}
    current_heading = "CONTACT"

    for raw_line in resume.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        normalized = re.sub(r"\s+", " ", line.upper())
        if normalized in known_headings:
            current_heading = "TECHNICAL SKILLS" if normalized == "LANGUAGES & SKILLS" else normalized
            sections.setdefault(current_heading, [current_heading])
            continue

        sections.setdefault(current_heading, [])
        sections[current_heading].append(line)

    return {heading: "\n".join(lines).strip() for heading, lines in sections.items()}


def build_ai_tailored_resume(resume: str, job_description: str) -> str:
    client = OpenAI(api_key=get_api_key())
    response = client.chat.completions.create(
        model=get_model(),
        temperature=0.2,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an expert resume editor. Create a complete tailored resume draft "
                    "for the target internship. Do not invent companies, dates, degrees, tools, "
                    "certifications, metrics, or experience. Only reorganize and rewrite what is "
                    "supported by the original resume. If a requested skill is missing, include it "
                    "only as a suggested project or learning gap, not as claimed experience."
                ),
            },
            {
                "role": "user",
                "content": f"Original Resume:\n{resume}\n\nJob Description:\n{job_description}",
            },
        ],
    )
    return response.choices[0].message.content or ""


def create_docx_download(resume_text: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(10)

    for block_index, block in enumerate(resume_text.split("\n\n")):
        lines = [line for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        if block_index == 0:
            for line_index, line in enumerate(lines):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                run = paragraph.add_run(line)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12 if line_index == 0 else 9)
                run.bold = line_index == 0
            continue

        if lines[0].isupper() and len(lines[0]) < 60:
            heading = document.add_paragraph()
            heading.paragraph_format.space_before = Pt(6)
            heading.paragraph_format.space_after = Pt(1)
            heading_run = heading.add_run(lines[0])
            heading_run.bold = True
            heading_run.font.name = "Times New Roman"
            heading_run.font.size = Pt(11)
            add_bottom_border(heading)
            lines = lines[1:]
        for line in lines:
            if line.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                run = paragraph.add_run(line[2:])
            else:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                run = paragraph.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9.5)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def create_preserved_docx_download(original_docx: bytes, result: dict[str, Any], job_description: str) -> bytes:
    document = Document(BytesIO(original_docx))
    strong_matches = result.get("strong_matches", [])
    demo_rewrite_bullets(document, job_description)

    if not strong_matches:
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    targeted_line = "Targeted Skills: " + ", ".join(strong_matches[:8])
    if any("targeted skills:" in paragraph.text.lower() for paragraph in document.paragraphs):
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    insert_after_index = find_skills_insert_index(document)
    if insert_after_index is None:
        paragraph = document.add_paragraph(targeted_line)
    else:
        source_paragraph = document.paragraphs[insert_after_index]
        paragraph = insert_paragraph_after(source_paragraph, targeted_line)
        copy_paragraph_style(source_paragraph, paragraph)

    if paragraph.runs:
        paragraph.runs[0].bold = False

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def demo_rewrite_bullets(document: Document, job_description: str) -> None:
    job_keywords = set(extract_keywords(job_description))
    replacements = 0

    for paragraph in document.paragraphs:
        if replacements >= 3:
            break
        if not is_bullet_paragraph(paragraph):
            continue

        original = paragraph.text.strip()
        rewritten = demo_rewrite_bullet(original, job_keywords)
        if rewritten and rewritten != original:
            replace_paragraph_text(paragraph, rewritten)
            replacements += 1


def is_bullet_paragraph(paragraph: Any) -> bool:
    style_name = (paragraph.style.name or "").lower()
    has_numbering = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    return "bullet" in style_name or has_numbering


def demo_rewrite_bullet(original: str, job_keywords: set[str]) -> str:
    text = original.lower()

    if "support" in text and any(keyword in job_keywords for keyword in {"it support", "troubleshooting", "communication"}):
        return (
            "Resolved daily hardware, software, and network support requests for multiple county departments; "
            "documented issues clearly and delivered onsite and remote assistance through the internal help-desk system."
        )

    if "workstation migration" in text or "laptop-docking" in text:
        return (
            "Supported a county-wide workstation migration for 50+ employees by assisting with hardware setup, "
            "software configuration, and user onboarding to improve mobility and workflow efficiency."
        )

    if "database" in text and any(keyword in job_keywords for keyword in {"sql", "data analysis", "documentation"}):
        return (
            "Designed a normalized SQL database schema for a multi-role e-commerce marketplace, supporting product listings, "
            "orders, reviews, helpdesk workflows, and reliable transactional data operations."
        )

    if "documentation" in text and "documentation" in job_keywords:
        return (
            "Updated technical documentation with clearer API usage examples, making onboarding and future maintenance easier "
            "for incoming developers."
        )

    return original


def replace_paragraph_text(paragraph: Any, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    first_run = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = text


def find_skills_insert_index(document: Document) -> int | None:
    skills_headings = {"technical skills", "skills", "languages & skills"}
    current_section = ""
    last_skills_paragraph_index = None

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        normalized = re.sub(r"\s+", " ", text.lower())
        if normalized in skills_headings:
            current_section = normalized
            last_skills_paragraph_index = index
            continue
        if current_section and text.isupper() and len(text) < 60:
            break
        if current_section:
            last_skills_paragraph_index = index

    return last_skills_paragraph_index


def insert_paragraph_after(paragraph: Any, text: str) -> Any:
    new_paragraph = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph)
    inserted = DocxParagraph(new_paragraph, paragraph._parent)
    inserted.add_run(text)
    return inserted


def copy_paragraph_style(source: Any, target: Any) -> None:
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing


def add_bottom_border(paragraph: Any) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    border.append(bottom)
    paragraph_properties.append(border)


def create_pdf_download(resume_text: str) -> bytes:
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=32,
        bottomMargin=32,
    )
    contact_name_style = ParagraphStyle(
        "ContactName",
        fontName="Times-Bold",
        fontSize=14,
        leading=15,
        alignment=TA_CENTER,
        spaceAfter=0,
    )
    contact_style = ParagraphStyle(
        "Contact",
        fontName="Times-Roman",
        fontSize=9,
        leading=10,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        fontName="Times-Bold",
        fontSize=10.5,
        leading=11,
        spaceBefore=6,
        spaceAfter=0,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        fontName="Times-Roman",
        fontSize=9.2,
        leading=10.2,
        spaceAfter=0,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=14,
        firstLineIndent=-7,
    )
    story = []

    for block_index, block in enumerate(resume_text.split("\n\n")):
        clean_block = block.strip()
        if not clean_block:
            continue
        lines = clean_block.splitlines()

        if block_index == 0:
            if lines:
                story.append(Paragraph(escape(lines[0]), contact_name_style))
            if len(lines) > 1:
                story.append(Paragraph(escape(" | ".join(lines[1:])), contact_style))
            continue

        if lines and lines[0].isupper() and len(lines[0]) < 60:
            story.append(Paragraph(escape(lines[0]), heading_style))
            story.append(HRFlowable(width="100%", thickness=0.6, color=colors.black, spaceBefore=0, spaceAfter=2))
            render_resume_lines(story, lines[1:], body_style, bullet_style)
        else:
            render_resume_lines(story, lines, body_style, bullet_style)
        story.append(Spacer(1, 4))

    document.build(story)
    return output.getvalue()


def render_resume_lines(story: list[Any], lines: list[str], body_style: ParagraphStyle, bullet_style: ParagraphStyle) -> None:
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
        if clean_line.startswith("- ") or clean_line[:1] in {"\u2022", "\u00b7"}:
            bullet_text = clean_line[2:].strip() if clean_line.startswith("- ") else clean_line[1:].strip()
            story.append(Paragraph(f"&bull; {escape(bullet_text)}", bullet_style))
        else:
            story.append(Paragraph(escape(clean_line), body_style))


def read_uploaded_resume(uploaded_file: Any) -> str:
    if uploaded_file is None:
        return ""

    uploaded_file.seek(0)
    raw = uploaded_file.read()
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        reader = PdfReader(BytesIO(raw))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()

    if file_name.endswith(".docx"):
        document = Document(BytesIO(raw))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(paragraphs).strip()

    if file_name.endswith(".txt"):
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("latin-1", errors="ignore")

    return ""


def get_uploaded_resume_bytes(uploaded_file: Any) -> bytes:
    if uploaded_file is None:
        return b""

    uploaded_file.seek(0)
    return uploaded_file.read()


def extract_job_description_from_url(url: str) -> str:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/126.0.0.0 Safari/537.36"
        )
    }
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "svg"]):
        tag.decompose()

    page_title = soup.title.get_text(" ", strip=True) if soup.title else ""
    meta_description = ""
    meta_tag = soup.find("meta", attrs={"name": "description"})
    if meta_tag and meta_tag.get("content"):
        meta_description = meta_tag["content"].strip()

    main_content = soup.find("main") or soup.find("article") or soup.body or soup
    text = main_content.get_text("\n", strip=True)
    lines = [line.strip() for line in text.splitlines() if len(line.strip()) > 2]
    cleaned_text = "\n".join(dict.fromkeys(lines))

    extracted = "\n\n".join(part for part in [page_title, meta_description, cleaned_text] if part)
    return extracted[:12000].strip()


def analyze_resume(resume: str, job_description: str) -> dict[str, Any]:
    client = OpenAI(api_key=get_api_key())

    response = client.chat.completions.create(
        model=get_model(),
        temperature=0.2,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    "Resume:\n"
                    f"{resume}\n\n"
                    "Job Description:\n"
                    f"{job_description}"
                ),
            },
        ],
    )

    content = response.choices[0].message.content or "{}"
    return json.loads(content)


def render_list(items: list[str], empty_text: str) -> None:
    if not items:
        st.write(empty_text)
        return

    for item in items:
        st.markdown(f"- {item}")


st.title("AI Internship Hunter")
st.caption("Resume-JD Match Analyzer")

api_key = get_api_key()
use_demo_mode = st.sidebar.toggle("Demo Mode", value=True)
if use_demo_mode:
    st.sidebar.info("Free keyword-based analysis. No API key needed.")
elif not api_key:
    st.warning("Add your OpenAI API key to continue, or turn on Demo Mode.")

st.subheader("1. Upload or Paste Resume")
uploaded_resume = st.file_uploader(
    "Upload resume",
    type=["pdf", "docx", "txt"],
    help="Upload a Word .docx for the best formatting preservation. PDF works for text extraction, but not perfect layout editing.",
)
st.info("For the cleanest tailored resume, upload the Word .docx version. PDF resumes can be read, but exact formatting is harder to preserve.")

resume_from_file = ""
if uploaded_resume is not None:
    try:
        if uploaded_resume.name.lower().endswith(".docx"):
            st.session_state.original_docx_bytes = get_uploaded_resume_bytes(uploaded_resume)
        else:
            st.session_state.original_docx_bytes = b""

        resume_from_file = read_uploaded_resume(uploaded_resume)
        if uploaded_resume.name.lower().endswith(".pdf"):
            st.warning("PDF uploaded: I can extract the text, but a Word .docx version is better for precise resume editing.")
        if not resume_from_file:
            st.warning("I could not extract text from this file. Try pasting the resume text below.")
    except Exception as exc:
        st.warning(f"I could not read this resume file: {exc}")

resume_text = st.text_area(
    "Resume",
    value=resume_from_file,
    height=220,
    placeholder="Paste your resume text here...",
)

st.subheader("2. Add Job Description")
job_url = st.text_input(
    "Job posting URL",
    placeholder="Paste a job posting link here...",
)

if "job_description_text" not in st.session_state:
    st.session_state.job_description_text = ""

extract_clicked = st.button("Extract Job Description from URL", use_container_width=True)
if extract_clicked:
    if not job_url.strip():
        st.warning("Please paste a job posting URL first.")
    else:
        with st.spinner("Extracting job description from the page..."):
            try:
                extracted_job_description = extract_job_description_from_url(job_url.strip())
            except Exception as exc:
                st.warning(
                    "I could not extract this page. Some job sites block automatic reading. "
                    "Please paste the job description manually below."
                )
            else:
                if extracted_job_description:
                    st.session_state.job_description_text = extracted_job_description
                    st.success("Job description extracted. You can edit it below before analyzing.")
                else:
                    st.warning("I could not find readable text on this page. Please paste the job description manually.")

job_description = st.text_area(
    "Job Description",
    value=st.session_state.job_description_text,
    height=220,
    placeholder="Paste the internship job description here...",
)
st.session_state.job_description_text = job_description

analyze_clicked = st.button("Analyze", type="primary", use_container_width=True)

if analyze_clicked:
    if not use_demo_mode and not api_key:
        st.error("Missing OpenAI API key. Turn on Demo Mode to analyze for free.")
    elif not resume_text.strip():
        st.error("Please upload or paste your resume first.")
    elif not job_description.strip():
        st.error("Please paste the job description first.")
    else:
        with st.spinner("Analyzing your resume match..."):
            try:
                if use_demo_mode:
                    result = demo_analyze_resume(resume_text, job_description)
                else:
                    result = analyze_resume(resume_text, job_description)
            except Exception as exc:
                error_text = str(exc)
                if "insufficient_quota" in error_text or "Error code: 429" in error_text:
                    st.error(
                        "Your OpenAI API key has no available quota. "
                        "Check your OpenAI billing/credits, then try again with an active API key."
                    )
                else:
                    st.error(f"Analysis failed: {exc}")
                st.stop()

        st.session_state.analysis_result = result
        st.session_state.analysis_resume = resume_text
        st.session_state.analysis_job_description = job_description
        st.session_state.analysis_original_docx_bytes = st.session_state.get("original_docx_bytes", b"")
        st.session_state.tailored_resume = ""

if st.session_state.get("analysis_result"):
    result = st.session_state.analysis_result
    score = int(result.get("match_score", 0))
    strong_matches = result.get("strong_matches", [])
    missing_keywords = result.get("missing_keywords", [])
    suggestions = result.get("resume_suggestions", [])
    tailored_bullets = result.get("tailored_bullets", [])

    st.divider()
    st.header(f"Match Score: {score}/100")
    st.progress(max(0, min(score, 100)) / 100)

    st.subheader("Strong Matches")
    render_list(strong_matches, "No strong matches found yet.")

    st.subheader("Missing Keywords")
    render_list(missing_keywords, "No major missing keywords found.")

    st.subheader("Resume Suggestions")
    render_list(suggestions, "No suggestions returned.")

    st.subheader("Tailored Bullet Points")
    render_list(tailored_bullets, "No tailored bullets returned.")

    st.subheader("Tailored Resume Draft")
    if st.button("Generate Tailored Resume Draft", use_container_width=True):
        with st.spinner("Creating a tailored resume draft..."):
            try:
                if use_demo_mode:
                    tailored_resume = build_demo_tailored_resume(
                        st.session_state.analysis_resume,
                        st.session_state.analysis_job_description,
                        result,
                    )
                else:
                    tailored_resume = build_ai_tailored_resume(
                        st.session_state.analysis_resume,
                        st.session_state.analysis_job_description,
                    )
            except Exception as exc:
                st.error(f"Could not generate tailored resume: {exc}")
                st.stop()

        st.session_state.tailored_resume = tailored_resume

    if st.session_state.get("tailored_resume"):
        st.text_area(
            "Preview",
            value=st.session_state.tailored_resume,
            height=360,
        )

        original_docx = st.session_state.get("analysis_original_docx_bytes", b"")
        if original_docx:
            docx_file = create_preserved_docx_download(
                original_docx,
                result,
                st.session_state.analysis_job_description,
            )
            st.caption("Word download preserves the uploaded .docx formatting and applies conservative targeted bullet rewrites.")
        else:
            docx_file = create_docx_download(st.session_state.tailored_resume)
        pdf_file = create_pdf_download(st.session_state.tailored_resume)
        left_column, right_column = st.columns(2)
        with left_column:
            st.download_button(
                "Download Word",
                data=docx_file,
                file_name="tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with right_column:
            st.download_button(
                "Download PDF",
                data=pdf_file,
                file_name="tailored_resume.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
